import io
import os
import re
import zipfile
import subprocess
import tempfile
import datetime
import requests
import streamlit as st
import streamlit.components.v1 as components
from docx import Document

# ==========================================
# PAGE CONFIGURATION
# ==========================================
st.set_page_config(
    page_title="Document Generator | Fidelity Funding",
    page_icon="📄",
    layout="wide"
)

# Injected directly into the parent window's head to prevent 'Enter' key from submitting forms
components.html("""
    <script>
    const parentDoc = window.parent.document;
    if (!parentDoc.getElementById('prevent-enter-submit')) {
        const script = parentDoc.createElement('script');
        script.id = 'prevent-enter-submit';
        script.type = 'text/javascript';
        script.innerHTML = `
            document.addEventListener('keydown', function(e) {
                // Target standard input fields and intercept the Enter key
                if ((e.key === 'Enter' || e.keyCode === 13) && e.target.tagName === 'INPUT') {
                    e.preventDefault();
                    e.stopPropagation();
                    e.stopImmediatePropagation(); // Crucial for overriding React's synthetic events
                }
            }, true); // 'true' means capture phase, catching it before React does
        `;
        parentDoc.head.appendChild(script);
    }
    </script>
""", height=0, width=0)

# ==========================================
# EASY CONFIGURATION / LINK MANAGEMENT
# ==========================================
TEMPLATE_CONFIG = {
    "RPS_CLASSES": {
        "RPS-Y | 1yr | 10.0%": {
            "class_code": "RPS-Y",
            "doc_url": "https://docs.google.com/document/d/1u5_-1qhJOk_6HfRILFRLTBAxL3UmE-p1/"
        },
        "RPS-Z | 1yr | 6.0%": {
            "class_code": "RPS-Z",
            "doc_url": "https://docs.google.com/document/d/1u5_-1qhJOk_6HfRILFRLTBAxL3UmE-p1/"
        }
    },
    "DEED_OF_ADHERENCE": "https://docs.google.com/document/d/1WhyRxt6nOaReWplabFUVUXgfgZI_zWtS/",
    "DECLARATION_FORM": "https://docs.google.com/document/d/1NBmaPDk5RwJZArVO4OUxWDI33q8Ox_U4/"
}

# ==========================================
# HELPER FUNCTIONS
# ==========================================

def extract_doc_id(url: str) -> str:
    """Extract Google Doc ID from full URL."""
    match = re.search(r"/d/([a-zA-Z0-9-_]+)", url)
    return match.group(1) if match else ""

def fetch_google_doc_bytes(url: str) -> bytes:
    """Download Google Doc directly as DOCX bytes."""
    doc_id = extract_doc_id(url)
    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=docx"
    response = requests.get(export_url)
    response.raise_for_status()
    return response.content

def replace_placeholders_in_paragraph(paragraph, replacements: dict):
    """Replace placeholder keys in paragraph while preserving text runs."""
    for key, value in replacements.items():
        if key in paragraph.text:
            # Replace at individual run level if key is contained in a single run
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)
            
            # If key spans across multiple runs, replace across paragraph text
            if key in paragraph.text:
                full_text = paragraph.text.replace(key, value)
                for i, run in enumerate(paragraph.runs):
                    if i == 0:
                        run.text = full_text
                    else:
                        run.text = ""

def process_docx_bytes(file_bytes: bytes, replacements: dict) -> bytes:
    """Load DOCX bytes, replace placeholders, and return modified DOCX bytes."""
    doc_io = io.BytesIO(file_bytes)
    doc = Document(doc_io)

    # Process all paragraphs
    for p in doc.paragraphs:
        replace_placeholders_in_paragraph(p, replacements)

    # Process all table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_in_paragraph(p, replacements)

    output_io = io.BytesIO()
    doc.save(output_io)
    return output_io.getvalue()

def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes to PDF bytes using headless LibreOffice."""
    with tempfile.TemporaryDirectory() as tmpdir:
        input_docx_path = os.path.join(tmpdir, "input.docx")
        with open(input_docx_path, "wb") as f:
            f.write(docx_bytes)
        
        cmd = [
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", tmpdir,
            input_docx_path
        ]
        
        try:
            subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
            output_pdf_path = os.path.join(tmpdir, "input.pdf")
            if os.path.exists(output_pdf_path):
                with open(output_pdf_path, "rb") as f:
                    return f.read()
        except Exception as e:
            st.warning(f"PDF conversion warning: {e}. Falling back to DOCX format.")
    return None

# ==========================================
# USER INTERFACE
# ==========================================

st.title("📄 PDF & Document Generator")
st.markdown("Automated generation of **Subscription Agreements**, **Deed of Adherence**, and **Declaration Forms**.")

# Sidebar for Config Management
with st.sidebar:
    st.header("⚙️ Template Link Management")
    st.info("You can quickly update Google Doc template links below if needed.")
    
    selected_rps = st.selectbox("Select RPS Class", list(TEMPLATE_CONFIG["RPS_CLASSES"].keys()))
    rps_sub_url = st.text_input("Subscription Agreement Link", TEMPLATE_CONFIG["RPS_CLASSES"][selected_rps]["doc_url"])
    deed_url = st.text_input("Deed of Adherence Link", TEMPLATE_CONFIG["DEED_OF_ADHERENCE"])
    decl_url = st.text_input("Declaration Form Link", TEMPLATE_CONFIG["DECLARATION_FORM"])

# Main Input Form
with st.form("doc_generation_form"):
    st.subheader("1. Client & Investment Details")
    col1, col2 = st.columns(2)
    
    with col1:
        client_name = st.text_input("Client Name", value="")
        client_nric = st.text_input("NRIC / Passport / Reg No", value="")
        client_email = st.text_input("Client Email", value="")
        client_contact = st.text_input("Contact Number", value="+60")
        client_address = st.text_area("Correspondence Address", value="")
        
    with col2:
        client_dob = st.text_input("Date of Birth", value="")
        client_nationality = st.text_input("Nationality", value="")
        client_occupation = st.text_input("Occupation", value="")
        investment_amt = st.text_input("Investment Amount (RM)", value="")
        agreement_date = st.date_input("Agreement Date", value=datetime.date.today())
        stamping = st.selectbox("Stamping", options=["Yes", "No"], index=0)

    st.subheader("2. Bank Details")
    b_col1, b_col2, b_col3 = st.columns(3)
    with b_col1:
        bank_name = st.text_input("Bank Name", value="")
    with b_col2:
        bank_acc_name = st.text_input("Bank Account Name", value="")
    with b_col3:
        bank_acc_no = st.text_input("Bank Account Number", value="")

    st.subheader("3. Witness / Advisor Information")
    w_col1, w_col2 = st.columns(2)
    with w_col1:
        witness_name = st.text_input("Witness / Advisor Name", value="")
    with w_col2:
        witness_nric = st.text_input("Witness / Advisor NRIC / Passport", value="")

    st.subheader("4. Nominee Information")
    
    nom_tabs = st.tabs(["Nominee 1", "Nominee 2", "Nominee 3", "Nominee 4"])
    nom_data = {}
    
    for i, tab in enumerate(nom_tabs, start=1):
        with tab:
            n_col1, n_col2 = st.columns(2)
            with n_col1:
                nom_data[f"<<NOM{i}_NAME>>"] = st.text_input(f"Nominee {i} Name", key=f"n_name_{i}")
                nom_data[f"<<NOM{i}_NRIC>>"] = st.text_input(f"Nominee {i} NRIC/Passport", key=f"n_nric_{i}")
                nom_data[f"<<NOM{i}_RELATIONSHIP>>"] = st.text_input(f"Nominee {i} Relationship", key=f"n_rel_{i}")
            with n_col2:
                nom_data[f"<<NOM{i}_ADDRESS>>"] = st.text_area(f"Nominee {i} Address", key=f"n_addr_{i}")
                nom_data[f"<<NOM{i}_EMAIL>>"] = st.text_input(f"Nominee {i} Email", key=f"n_email_{i}")
                nom_data[f"<<NOM{i}_PERCENTAGE>>"] = st.text_input(f"Nominee {i} Percentage (%)", key=f"n_pct_{i}")

    submit_button = st.form_submit_button("🔨 Generate Documents & PDF")

# ==========================================
# VALIDATION & PROCESSING
# ==========================================

if submit_button:
    errors = []

    # 1. Mandatory Client Name Check
    if not client_name.strip():
        errors.append("Please enter the Client Name before proceeding.")

    # 2. Client Email Validation
    if client_email.strip():
        email_regex = r"^[\w\.-]+@[\w\.-]+\.\w+$"
        if not re.match(email_regex, client_email.strip()):
            errors.append("Invalid Client Email format.")

    # 3. Investment Amount Validation (Numbers only)
    clean_inv_amt = investment_amt.strip().replace(",", "")
    if clean_inv_amt and not re.match(r"^\d+(\.\d+)?$", clean_inv_amt):
        errors.append("Investment Amount must contain numbers only.")

    # 4. Nominee Percentage Validation & 100% Sum Check
    total_pct = 0.0
    pct_has_error = False

    for i in range(1, 5):
        pct_val_str = nom_data.get(f"<<NOM{i}_PERCENTAGE>>", "").strip()
        if pct_val_str:
            try:
                pct_val = float(pct_val_str)
                total_pct += pct_val
            except ValueError:
                errors.append(f"Nominee {i} Percentage must be a valid number.")
                pct_has_error = True

    if not pct_has_error and abs(total_pct - 100.0) > 0.001:
        errors.append(f"The sum of all 4 nominee percentages must equal 100%. (Current total: {total_pct:.2f}%)")

    # Display Errors or Proceed
    if errors:
        for err in errors:
            st.error(f"⚠️ {err}")
    else:
        with st.spinner("Fetching templates and generating PDFs..."):
            formatted_date = agreement_date.strftime("%d/%m/%Y")

            # Map input fields to placeholders
            raw_replacements = {
                "<<CLIENT_NAME>>": client_name,
                "<<CLIENT_NRIC>>": client_nric,
                "<<CLIENT_EMAIL>>": client_email,
                "<<CLIENT_CONTACT>>": client_contact,
                "<<CLIENT_ADDRESS>>": client_address,
                "<<CLIENT_DOB>>": client_dob,
                "<<CLIENT_NATIONALITY>>": client_nationality,
                "<<CLIENT_OCCUPATION>>": client_occupation,
                "<<INVESTMENT_AMT>>": investment_amt,
                "<<DATE>>": formatted_date,
                "<<STAMPING>>": stamping,
                "<<BANK_NAME>>": bank_name,
                "<<BANK_ACC_NAME>>": bank_acc_name,
                "<<BANK_ACC_NO>>": bank_acc_no,
                "<<WITNESS_NAME>>": witness_name,
                "<<WITNESS_NRIC>>": witness_nric,
                "<<RPS-CLASS>>": TEMPLATE_CONFIG["RPS_CLASSES"][selected_rps]["class_code"],
                **nom_data
            }

            # If any field is left blank, replace placeholder with two spaces "  "
            replacements = {k: (v.strip() if v and str(v).strip() else "  ") for k, v in raw_replacements.items()}

            clean_client_name = client_name.strip()
            clean_date = agreement_date.strftime("%Y%m%d")
            rps_code = TEMPLATE_CONFIG["RPS_CLASSES"][selected_rps]["class_code"]

            # Output File Naming Patterns
            fn_sub_base = f"{clean_client_name} 1. FF {rps_code} SubscriptionAgreement {clean_date}"
            fn_deed_base = f"{clean_client_name} 2. FF Deed of Adherence {clean_date}"
            fn_decl_base = f"{clean_client_name} 3. FF Declaration Form (Sophisticated Investor)"

            try:
                # 1. Fetch Google Doc templates
                sub_docx_raw = fetch_google_doc_bytes(rps_sub_url)
                deed_docx_raw = fetch_google_doc_bytes(deed_url)
                decl_docx_raw = fetch_google_doc_bytes(decl_url)

                # 2. Populate placeholders
                sub_docx = process_docx_bytes(sub_docx_raw, replacements)
                deed_docx = process_docx_bytes(deed_docx_raw, replacements)
                decl_docx = process_docx_bytes(decl_docx_raw, replacements)

                # 3. Convert to PDF
                sub_pdf = convert_docx_to_pdf(sub_docx)
                deed_pdf = convert_docx_to_pdf(deed_docx)
                decl_pdf = convert_docx_to_pdf(decl_docx)

                st.success("✅ Documents generated successfully!")

                # Prepare ZIP download package
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                    if sub_pdf:
                        zip_file.writestr(f"{fn_sub_base}.pdf", sub_pdf)
                        zip_file.writestr(f"{fn_deed_base}.pdf", deed_pdf)
                        zip_file.writestr(f"{fn_decl_base}.pdf", decl_pdf)
                    # Add DOCX versions as well
                    zip_file.writestr(f"{fn_sub_base}.docx", sub_docx)
                    zip_file.writestr(f"{fn_deed_base}.docx", deed_docx)
                    zip_file.writestr(f"{fn_decl_base}.docx", decl_docx)

                # Single Download Button for ZIP
                st.download_button(
                    label="📦 Download All Documents (ZIP)",
                    data=zip_buffer.getvalue(),
                    file_name=f"{clean_client_name}_Documents.zip",
                    mime="application/zip",
                    type="primary"
                )

                st.markdown("---")
                st.subheader("📥 Individual File Downloads")

                col_d1, col_d2, col_d3 = st.columns(3)

                with col_d1:
                    st.write("**1. Subscription Agreement**")
                    if sub_pdf:
                        st.download_button("Download PDF", sub_pdf, f"{fn_sub_base}.pdf", "application/pdf")
                    st.download_button("Download DOCX", sub_docx, f"{fn_sub_base}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                with col_d2:
                    st.write("**2. Deed of Adherence**")
                    if deed_pdf:
                        st.download_button("Download PDF", deed_pdf, f"{fn_deed_base}.pdf", "application/pdf")
                    st.download_button("Download DOCX", deed_docx, f"{fn_deed_base}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                with col_d3:
                    st.write("**3. Declaration Form**")
                    if decl_pdf:
                        st.download_button("Download PDF", decl_pdf, f"{fn_decl_base}.pdf", "application/pdf")
                    st.download_button("Download DOCX", decl_docx, f"{fn_decl_base}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            except Exception as e:
                st.error(f"Error processing Google Docs: {e}")
